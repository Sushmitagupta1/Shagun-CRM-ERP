import io
import os

import pytest

from docx import Document

from app.services.word_export import build_menu_docx


def _tiny_png(tmp_path) -> str:
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (4, 4), (200, 180, 120)).save(buf, format="PNG")
    path = os.path.join(str(tmp_path), "bg.png")
    with open(path, "wb") as f:
        f.write(buf.getvalue())
    return path


def test_build_menu_docx_with_background(tmp_path):
    img = _tiny_png(tmp_path)
    data = build_menu_docx(
        [
            {"text": "STARTERS", "is_heading": True, "page": 0},
            {"text": "Paneer Tikka", "is_heading": False, "page": 0},
        ],
        img,
        {"heading": "#5A0016", "item": "#8C6A1F", "desc": "#4B5563"},
    )
    doc = Document(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    assert "STARTERS" in texts
    assert "Paneer Tikka" in texts
    head = next(p for p in doc.paragraphs if p.text == "STARTERS")
    assert head.runs[0].bold is True
    assert head.runs[0].font.size.pt == 16
    assert str(head.runs[0].font.color.rgb) == "5A0016"
    header = doc.sections[0].header
    assert len(header.paragraphs[0].runs) == 1
    assert "<w:drawing>" in header.paragraphs[0].runs[0]._element.xml


def test_build_menu_docx_without_background():
    data = build_menu_docx(
        [{"text": "Paneer", "is_heading": False, "page": 0}],
        None,
        {"heading": "#5A0016", "item": "#8C6A1F", "desc": "#4B5563"},
    )
    doc = Document(io.BytesIO(data))
    assert doc.paragraphs[0].text == "Paneer"
    assert len(doc.sections[0].header.paragraphs[0].runs) == 0


def test_build_menu_docx_page_breaks():
    data = build_menu_docx(
        [
            {"text": "STARTERS", "is_heading": True, "page": 0},
            {"text": "Paneer Tikka", "is_heading": False, "page": 0},
            {"text": "DESSERTS", "is_heading": True, "page": 1},
        ],
        None,
        {"heading": "#5A0016", "item": "#8C6A1F", "desc": "#4B5563"},
    )
    doc = Document(io.BytesIO(data))
    assert 'w:type="page"' in doc.element.xml


def test_build_menu_docx_bad_hex_color_falls_back():
    data = build_menu_docx(
        [{"text": "Paneer", "is_heading": True, "page": 0}],
        None,
        {"heading": "not-a-color", "item": "#8C6A1F", "desc": "#4B5563"},
    )
    doc = Document(io.BytesIO(data))
    assert str(doc.paragraphs[0].runs[0].font.color.rgb) == "5A0016"


def test_build_menu_docx_missing_template_path_ignored():
    data = build_menu_docx(
        [{"text": "Paneer", "is_heading": False, "page": 0}],
        "C:/does/not/exist/bg.png",
        {"heading": "#5A0016", "item": "#8C6A1F", "desc": "#4B5563"},
    )
    doc = Document(io.BytesIO(data))
    assert doc.paragraphs[0].text == "Paneer"


def _first_template():
    base = os.path.join(os.path.dirname(__file__), "..", "templates")
    if not os.path.isdir(base):
        return None
    for cat in sorted(os.listdir(base)):
        cat_dir = os.path.join(base, cat)
        if not os.path.isdir(cat_dir):
            continue
        files = sorted(
            f for f in os.listdir(cat_dir)
            if f.lower().endswith((".jpg", ".jpeg", ".png", ".webp", ".jfif"))
        )
        if files:
            return os.path.join(cat_dir, files[0])
    return None


def test_build_menu_docx_with_real_template_image():
    """Real template images (incl. EXIF-heavy .jfif and .webp) must not break
    the export — python-docx cannot parse some of them directly, so the builder
    re-encodes the image via PIL first."""
    path = _first_template()
    if not path:
        pytest.skip("No template image available for the export test")
    data = build_menu_docx(
        [{"text": "STARTERS", "is_heading": True, "page": 0}],
        path,
        {"heading": "#5A0016", "item": "#8C6A1F", "desc": "#4B5563"},
    )
    doc = Document(io.BytesIO(data))
    assert doc.paragraphs[0].text == "STARTERS"
    assert len(doc.sections[0].header.paragraphs[0].runs) == 1
