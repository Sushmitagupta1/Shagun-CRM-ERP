import io
import os
import zipfile

import pytest

from docx import Document


def build_docx(paragraphs: list[tuple[str, list[str]]]) -> bytes:
    body = []
    for ppr, runs in paragraphs:
        ppr_xml = f"<w:pPr>{ppr}</w:pPr>" if ppr else ""
        body.append(f"<w:p>{ppr_xml}{''.join(runs)}</w:p>")
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"<w:body>{''.join(body)}</w:body></w:document>"
    )
    styles_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/></w:style>'
        "<w:style w:type=\"paragraph\" w:styleId=\"Heading1\"><w:name w:val=\"heading 1\"/></w:style>"
        "</w:styles>"
    )
    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '<Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
        "</Types>"
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        "</Relationships>"
    )
    doc_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>'
        "</Relationships>"
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml)
        zf.writestr("word/_rels/document.xml.rels", doc_rels)
        zf.writestr("word/styles.xml", styles_xml)
    return buf.getvalue()


def _first_template():
    base = os.path.join(os.path.dirname(__file__), "..", "templates")
    if not os.path.isdir(base):
        return None, None
    for cat in sorted(os.listdir(base)):
        cat_dir = os.path.join(base, cat)
        if not os.path.isdir(cat_dir):
            continue
        files = sorted(
            f for f in os.listdir(cat_dir)
            if f.lower().endswith((".jpg", ".jpeg", ".png", ".webp", ".jfif"))
        )
        if files:
            return cat, files[0]
    return None, None


async def login(client, username, password):
    resp = await client.post("/api/auth/login", json={"username": username, "password": password})
    assert resp.status_code == 200, f"Login failed: {resp.status_code} {resp.text}"
    return resp.json()["access_token"]


def auth(token):
    return {"Authorization": f"Bearer {token}"}


async def test_parse_word_docx_returns_lines(client):
    token = await login(client, "admin@shaguncatering.com", "admin123")
    docx = build_docx([
        ("", ['<w:r><w:rPr><w:b/></w:rPr><w:t xml:space="preserve">STARTERS</w:t></w:r>']),
        ("", ['<w:r><w:t xml:space="preserve">Paneer Tikka</w:t></w:r>']),
    ])
    resp = await client.post(
        "/api/menu/parse-word",
        headers=auth(token),
        files={"file": ("menu.docx", docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert data["file_name"] == "menu.docx"
    assert data["lines"] == [
        {"text": "STARTERS", "is_heading": True, "page": 0},
        {"text": "Paneer Tikka", "is_heading": False, "page": 0},
    ]


async def test_parse_word_rejects_unsupported_file(client):
    token = await login(client, "admin@shaguncatering.com", "admin123")
    resp = await client.post("/api/menu/parse-word", headers=auth(token), files={"file": ("menu.txt", b"hello", "text/plain")})
    assert resp.status_code == 400


async def test_parse_word_requires_auth(client):
    resp = await client.post("/api/menu/parse-word", files={"file": ("menu.docx", b"x", "application/octet-stream")})
    assert resp.status_code in (401, 403)


async def test_export_word_returns_docx(client):
    token = await login(client, "admin@shaguncatering.com", "admin123")
    cat, file = _first_template()
    if not file:
        pytest.skip("No template image available for the export test")
    resp = await client.post("/api/menu/export-word", headers=auth(token), json={
        "lines": [
            {"text": "STARTERS", "is_heading": True, "page": 0},
            {"text": "Paneer Tikka", "is_heading": False, "page": 0},
        ],
        "template_category": cat,
        "template_file": file,
        "colors": {"heading": "#5A0016", "item": "#8C6A1F", "desc": "#4B5563"},
    })
    assert resp.status_code == 200, resp.text
    assert "vnd.openxmlformats-officedocument.wordprocessingml" in resp.headers["content-type"]
    doc = Document(io.BytesIO(resp.content))
    texts = [p.text for p in doc.paragraphs]
    assert "STARTERS" in texts
    assert "Paneer Tikka" in texts
    head = next(p for p in doc.paragraphs if p.text == "STARTERS")
    assert head.runs[0].bold is True
    assert len(doc.sections[0].header.paragraphs[0].runs) == 1


async def test_export_word_empty_lines_rejected(client):
    token = await login(client, "admin@shaguncatering.com", "admin123")
    resp = await client.post("/api/menu/export-word", headers=auth(token), json={
        "lines": [],
        "template_category": "Wedding",
        "template_file": "x.jpg",
        "colors": {"heading": "#5A0016", "item": "#8C6A1F", "desc": "#4B5563"},
    })
    assert resp.status_code == 400
