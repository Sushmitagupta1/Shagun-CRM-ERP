import io
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Mm, Pt, RGBColor

PAGE_WIDTH = Mm(210)
PAGE_HEIGHT = Mm(297)


def _hex_color(value: str | None, default: str) -> RGBColor:
    v = (value or "").strip().lstrip("#")
    if len(v) == 6:
        try:
            return RGBColor(int(v[0:2], 16), int(v[2:4], 16), int(v[4:6], 16))
        except ValueError:
            pass
    return RGBColor(int(default[1:3], 16), int(default[3:5], 16), int(default[5:7], 16))


def _load_template_image(template_path: str) -> io.BytesIO:
    """Re-encode the template to a clean PNG so python-docx can always embed
    it. python-docx cannot parse some real-world images (EXIF-heavy .jfif files
    from WhatsApp, .webp, etc.), so we normalize them through PIL first."""
    from PIL import Image

    buf = io.BytesIO()
    with Image.open(template_path) as img:
        img.convert("RGB").save(buf, format="PNG")
    buf.seek(0)
    return buf


def build_menu_docx(lines: list[dict], template_path: str | None, colors: dict) -> bytes:
    """Build a .docx: template image as full-page background (in the header,
    behind the body text) + centered menu text with template-matched colours."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)
    section.header_distance = Mm(0)
    section.footer_distance = Mm(0)

    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(12)

    if template_path and os.path.isfile(template_path):
        header = section.header
        header.is_linked_to_previous = False
        hpara = header.paragraphs[0]
        hpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hpara.paragraph_format.space_before = Pt(0)
        hpara.paragraph_format.space_after = Pt(0)
        hpara.add_run().add_picture(_load_template_image(template_path), width=PAGE_WIDTH, height=PAGE_HEIGHT)

    heading_color = _hex_color(colors.get("heading"), "#5A0016")
    item_color = _hex_color(colors.get("item"), "#8C6A1F")

    last_page = 0
    for i, line in enumerate(lines):
        if i > 0 and line.get("page", 0) > last_page:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        last_page = line.get("page", 0)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.left_indent = Mm(18)
        para.paragraph_format.right_indent = Mm(18)
        run = para.add_run(line.get("text", ""))
        if line.get("is_heading"):
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = heading_color
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after = Pt(6)
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = item_color
            para.paragraph_format.space_after = Pt(4)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
